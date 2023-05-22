from docx import Document
import pandas as pd
import re

# Ask the user for the name of the document
filename = input("Please enter the name of the document: ")

# Load the document
doc = Document(filename)

# Initialize an empty dictionary to hold the acronyms and definitions
acronyms = {}

# Define regular expressions to match "phrase (ACRONYM)" and "ACRONYM (phrase)" patterns
# These patterns now match phrases that include words with hyphens and numerals
acronym_pattern_1 = re.compile(r'([A-Za-z0-9\- ]+)\s\(([A-Z0-9]{2,})\)')
acronym_pattern_2 = re.compile(r'\b([A-Z0-9]{2,})\s\(([A-Za-z0-9\- ]+)\)')

# Function to find acronyms in a given text
def find_acronyms(text):
    # Find all "phrase (ACRONYM)" matches in the text
    matches = re.findall(acronym_pattern_1, text)
    for match in matches:
        # Add the acronym and definition to our dictionary
        acronyms[match[1]] = match[0]
    
    # Find all "ACRONYM (phrase)" matches in the text
    matches = re.findall(acronym_pattern_2, text)
    for match in matches:
        # Add the acronym and definition to our dictionary
        acronyms[match[0]] = match[1]

# Loop through all paragraphs in the document
for para in doc.paragraphs:
    find_acronyms(para.text)

# Loop through all tables in the document
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            find_acronyms(cell.text)

# Put the acronyms and their definitions into a dataframe
df = pd.DataFrame(list(acronyms.items()), columns=['Acronym', 'Definition'])

# Save the dataframe to a CSV file
df.to_csv('acronyms.csv', index=False)
